#!/usr/bin/env python3
"""Проба скачивания статьи Дзен (диагностика + опционально docx).

Важно:
- URL вида https://dzen.ru/id/XXXX — это КАНАЛ, не статья.
- Статья: https://dzen.ru/a/agn86z3LoXbDuwvM
- requests без cookies → редирект SSO, контента нет (HTML ~3 KB, 0 тегов <p>).
"""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path
from urllib.parse import urlparse

import requests
from bs4 import BeautifulSoup

ROOT = Path(__file__).resolve().parents[1]
DEFAULT_OUT = ROOT / "03experiments" / "dzen_scrape"

HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
    ),
    "Accept-Language": "ru-RU,ru;q=0.9",
}


def classify_url(url: str) -> str:
    path = urlparse(url).path.strip("/")
    if path.startswith("a/"):
        return "article"
    if path.startswith("id/"):
        return "channel"
    return "unknown"


def is_sso_stub(html: str) -> bool:
    return "sso.dzen.ru" in html or "sso.passport.yandex" in html or len(html) < 8000


def fetch_html(url: str) -> tuple[str, int]:
    r = requests.get(url, headers=HEADERS, timeout=30)
    r.raise_for_status()
    return r.text, r.status_code


def extract_with_bs4(html: str) -> tuple[list[str], str | None]:
    soup = BeautifulSoup(html, "html.parser")
    selectors = [
        ("div.article-content__root", soup.select_one("div.article-content__root")),
        ("main", soup.find("main")),
        ("article", soup.find("article")),
        ("h1", soup.find("h1")),
    ]
    for name, node in selectors:
        if node and node.get_text(strip=True):
            return [node.get_text("\n", strip=True)], name

    paragraphs = [p.get_text(strip=True) for p in soup.find_all("p") if p.get_text(strip=True)]
    if paragraphs:
        return paragraphs, "p-tags"
    return [], None


def save_docx(paragraphs: list[str], title: str, path: Path) -> None:
    from docx import Document

    doc = Document()
    if title:
        doc.add_heading(title, level=1)
    for text in paragraphs:
        doc.add_paragraph(text)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def diagnose(url: str) -> int:
    kind = classify_url(url)
    print(f"URL: {url}")
    print(f"Тип: {kind} ({'статья' if kind == 'article' else 'канал' if kind == 'channel' else '?'})")

    html, status = fetch_html(url)
    print(f"HTTP {status}, размер HTML: {len(html)} байт")

    if is_sso_stub(html):
        print("\n[FAIL] Dzen returned SSO/antibot stub, not article HTML.")
        print("   requests + BeautifulSoup без сессии браузера не работают.")
        print("   Нужен: Playwright/Selenium с вашим логином ИЛИ ручное копирование.")
        return 2

    paragraphs, via = extract_with_bs4(html)
    print(f"Селектор: {via or 'не найден'}, блоков: {len(paragraphs)}")
    if not paragraphs:
        print("[FAIL] Content not extracted — check selectors in DevTools.")
        return 1
    print("[OK] Text found (preview):")
    print(paragraphs[0][:300])
    return 0


def main() -> int:
    parser = argparse.ArgumentParser(description="Проба скачивания статьи Дзен")
    parser.add_argument("url", help="URL статьи (/a/...) или канала (/id/...)")
    parser.add_argument("-o", "--output", type=Path, help="Путь к .docx (если удалось извлечь)")
    parser.add_argument("--diagnose-only", action="store_true", default=True)
    args = parser.parse_args()

    if classify_url(args.url) == "channel":
        print("[WARN] Channel URL, not a single publication.")
        print("   Open article -> Share -> Copy link (format /a/...).")
        print("   Пример статьи с вашего канала:")
        print("   https://dzen.ru/a/agn86z3LoXbDuwvM")
        print()

    code = diagnose(args.url)
    if code != 0 or args.diagnose_only or not args.output:
        return code

    html, _ = fetch_html(args.url)
    paragraphs, _ = extract_with_bs4(html)
    title = BeautifulSoup(html, "html.parser").find("title")
    save_docx(paragraphs, title.get_text() if title else "", args.output)
    print(f"Сохранено: {args.output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
